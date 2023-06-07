#pip install requests
#pip install python-docx

import requests
from docx import Document

def export_confluence_page(space_key, page_title, export_file_path, confluence_url, username, password):
    # Authenticate and get the page content from Confluence
    auth = (username, password)
    headers = {"Accept": "application/json"}
    url = "{}/rest/api/content".format(confluence_url)
    params = {
        "spaceKey": space_key,
        "title": page_title,
        "expand": "body.storage"
    }
    response = requests.get(url, params=params, auth=auth, headers=headers)
    response_json = response.json()

    if "results" not in response_json or len(response_json["results"]) == 0:
        print("Page not found")
        return

    page_id = response_json["results"][0]["id"]
    page_version = response_json["results"][0]["version"]["number"]
    page_content = response_json["results"][0]["body"]["storage"]["value"]

    # Create a Word document
    doc = Document()
    doc.add_heading(page_title, level=1)

    # Convert the Confluence markup to Word document format
    # Note: This conversion may not be perfect, and you might need to adjust the formatting as needed.
    # You can explore more advanced parsing libraries for better conversion results.
    # For simplicity, this example only handles basic formatting.
    # You might need to handle additional markup or specific cases depending on your Confluence content.
    paragraphs = page_content.split("\n")
    for paragraph in paragraphs:
        if paragraph.startswith("h1."):
            doc.add_heading(paragraph[4:], level=1)
        elif paragraph.startswith("h2."):
            doc.add_heading(paragraph[4:], level=2)
        elif paragraph.startswith("h3."):
            doc.add_heading(paragraph[4:], level=3)
        elif paragraph.startswith("h4."):
            doc.add_heading(paragraph[4:], level=4)
        elif paragraph.startswith("h5."):
            doc.add_heading(paragraph[4:], level=5)
        elif paragraph.startswith("h6."):
            doc.add_heading(paragraph[4:], level=6)
        else:
            doc.add_paragraph(paragraph)

    # Save the Word document
    doc.save(export_file_path)

    print("Export complete. File saved at: {}".format(export_file_path))


# Example usage
confluence_url = "https://your-confluence-instance.com"
space_key = "YOUR_SPACE_KEY"
page_title = "Your Page Title"
export_file_path = "path/to/exported_document.docx"
username = "your-username"
password = "your-password"

export_confluence_page(space_key, page_title, export_file_path, confluence_url, username, password)
